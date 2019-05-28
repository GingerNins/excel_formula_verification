from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.table import Table
from utils import add_field, shade_cell
from variable import Formula

__MINIMUM_ROWS = 7
__NUM_COLUMNS = 7
__COLUMN_WIDTHS = (0.37, 1.95, 0.69, 2.25, 0.77, 0.77, 0.77)
__TABLE_STYLE = 'Table Grid'
__ROW_HEIGHT = 0.1
__FILL_COLOR = 'd0cece'
__CELL_MARGINS = {'top': "50", 'bottom': "10", 'start': "50", 'end': "50"}
__AUTONUM = r'AUTONUM \s :'  # Document field code


test_num = 1  # FIXME: probs don't need this anymore


def add_table(doc: Document, formula: Formula):
    """
    Main function to add and setup a table for the verification document

    # NOTE: May end up having different types of tables, keep that in mind

    :param doc: Document where table will be added
    :param formula: Formula to base the document on
    :return: n/a
    """
    doc.add_paragraph()

    total_rows: int = __MINIMUM_ROWS + len(formula.variables)
    table: Table = doc.add_table(total_rows, __NUM_COLUMNS)
    table.style = __TABLE_STYLE
    __set_row_height(table)
    __set_column_widths(table)
    __set_margins(table)
    __create_headers(table, formula)
    __create_manual_formula_row(table, formula)
    __create_excel_formula_row(table, formula)
    __create_variable_rows(table, formula, total_rows)


def __set_column_widths(table: Table):
    """
    Sets the widths of the seven columns in the formula table
    :param table: Table to format
    :return: n/a
    """
    columns = table.columns

    for i, c in enumerate(columns):
        for cell in c.cells:
            cell.width = Inches(__COLUMN_WIDTHS[i])


def __set_row_height(table: Table):
    """
    Sets the height of the rows in the table
    :param table: Table to format
    :return: n/a
    """
    for r in table.rows:
        r.height = Inches(__ROW_HEIGHT)


def __set_margins(table: Table):
    """
    Sets the margins for each cell in the table
    by modifying the document's xml directly since
    docx package does not have this built-in
    :param table: Table to format
    :return: n/a
    """
    for c in table.columns:
        for cell in c.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')

            for k, v in __CELL_MARGINS.items():
                node = OxmlElement(f"w:{k}")
                node.set(qn('w:w'), v)
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)


def __create_headers(table: Table, formula: Formula):
    """
    Creates the header for the table, adds an auto-numbering field to the first cell to track the test number,
    adds the sheet name the formula is located on in the second cell.  Header cells are shaded.  Docx does not
    support cell shading, therefore the document xml is directly modified.
    :param table: table to add headers
    :param formula: formula to base the table on
    :return: n/a
    """
    table.cell(0, 0).merge(table.cell(0, 3))
    table.cell(0, 4).merge(table.cell(0, 6))

    cell = table.cell(0, 0)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    cell.shade_cell(__FILL_COLOR)

    paragraph = cell.paragraphs[0]
    run = paragraph.add_run('Test No. ')
    run.add_field(__AUTONUM)
    run.add_text(f" Name of test")  # TODO: Determine what the name of the test will be
    paragraph.style = 'Cell Heading'  # TODO: Create style that makes text based on a heading style so TOC can be made

    cell = table.cell(0, 6)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    cell.shade_cell(__FILL_COLOR)

    paragraph = cell.paragraphs[0]
    paragraph.text = f"Sheet: {formula.sheet}"
    paragraph.style = 'Cell Header Right'


def __create_manual_formula_row(table: Table, formula: Formula):
    """
    Creates a row containing a human-readable formula based on the Excel formula
    :param table: table to add formula
    :param formula: formula to add
    :return: n/a
    """
    cell = table.cell(1, 0)
    cell.text = '1'
    cell.paragraphs[0].style = 'Cell Text'
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # TODO: Add the formula and text
    # TODO: use LaTeX formatting from matplotlib to create the written formula
    # TODO: Excel document will need a column to write the LaTeX markup

    cell = table.cell(1, 1).merge(table.cell(1, 6))
    cell.text = 'INSERT MANUAL FORMULA'


def __create_excel_formula_row(table: Table, formula: Formula):
    """
    Creates a row containing the Excel readable formula to compare to human formula
    Cell shading is done using direct document xml modification
    :param table: table to update
    :param formula: formula to base the row on
    :return: n/a
    """
    table.cell(2, 0).merge(table.cell(3, 0))
    table.cell(2, 1).merge(table.cell(3, 5))

    cell = table.cell(2, 0)
    cell.text = '2'
    cell.paragraphs[0].style = 'Cell Text'
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    cell = table.cell(2, 6)
    cell.text = 'Pass?'
    cell.paragraphs[0].style = 'Cell Header Center'
    cell.shade_cell(__FILL_COLOR)

    cell = table.cell(2, 1)
    formula = '=' + formula.value
    cell.text = \
        'Compare the above formula to the Excel formula to determine if it should calculate the same as above:'
    cell.paragraphs[0].style = 'Cell Text'
    cell.add_paragraph(f'\n{formula}', 'Cell Text Center')


def __create_variable_rows(table: Table, formula: Formula, total_rows: int):
    # Merge as appropriate
    table.cell(4, 0).merge(table.cell(total_rows - 1, 0))
    table.cell(4, 4).merge(table.cell(4, 6))
    table.cell(total_rows - 2, 1).merge(table.cell(total_rows - 1, 1))
    table.cell(total_rows - 2, 2).merge(table.cell(total_rows - 1, 2))
    table.cell(total_rows - 2, 3).merge(table.cell(total_rows - 1, 3))

    # Numbered cell
    cell = table.cell(4, 0)
    cell.text = '3'
    cell.paragraphs[0].style = 'Cell Text'
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # Header row
    headers = ['Formula Variable', 'Cell', 'Range Name', 'Value']
    for i in range(1, 4+1):
        cell = table.cell(4, i)
        cell.text = headers[i - 1]
        cell.paragraphs[0].style = 'Cell Header'
        cell.shade_cell(__FILL_COLOR)

    # Individual Variable rows
    for i, var in enumerate(formula.variables):
        if var is None:
            continue
        table.cell(5 + i, 4).merge(table.cell(5 + i, 6))
        cell = table.cell(5 + i, 2)
        cell.text = var.coordinate or 'N/A'
        cell.paragraphs[0].style = 'Cell Text'

        cell = table.cell(5 + i, 3)
        cell.text = var.name
        cell.paragraphs[0].style = 'Cell Text'

        # TODO: work on figuring out what the value is supposed to be -- input/output?
        # TODO: May have to be manually entered?  Possibly in the spreadsheet
        # TODO: Two step process of creating the file containing the formulas, then entering the values they should be

        cell = table.cell(5 + i, 4)
        # NOTE: Possibilities for output values
        #  - Name/Formula/Variable (i: val/form, o: val)
        #  - Name/Formula/Variable (i: val/form, o: None) -
        #  - Name/Formula/Variable (i: None, o: None) - Blank cells?
        #  - Global Names?
        cell.text = var.output or ''  # FIXME: output not outputting, 0's causing issue?
        cell.paragraphs[0].style = 'Cell Text'

    # TODO: Add formula variable itself (Last Row)

    # TODO: Make separate function for this
    # Sub header rows
    headers = ['Manual', 'Excel', 'Pass']
    for i in range(4, 6+1):
        cell = table.cell(total_rows - 2, i)
        cell.text = headers[i - 4]
        cell.paragraphs[0].style = 'Cell Header Center'
        cell.shade_cell(__FILL_COLOR)
