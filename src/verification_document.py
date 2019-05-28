"""

"""
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from table import add_table
from utils import add_field, add_outline_level

# TODO: Insert a table of contents

# Document Field Codes
__PAGE: str = r'PAGE \* Arabic \* MERGEFORMAT'
__NUMPAGES: str = r'NUMPAGES \* Arabic \* MERGEFORMAT'


doc = docx.Document()


def create_document(template_data: dict, template_name: str, template_description: str,
                    margins: tuple = (0.5, 0.5, 0.5, 0.5),
                    tab_stops: tuple = (4.0, 7.5)) -> Document:
    """
    Main function to set up the verification test document
    :param template_data: dict containing lists: formulas, constants, names
    :param template_name: Name of the excel template
    :param template_description: Description of the excel template
    :param margins: Document margins
    :param tab_stops: Header/footer tab stops
    :return: Verification test document
    """

    __create_styles(tab_stops)
    __document_setup(template_name, template_description, margins)

    for f in template_data['formulas']:
        add_table(doc, f)

    # TODO: Process the constants and names if necessary

    return doc


def __create_styles(tab_stops: tuple):
    """
    Sets up all the styles for the document
    Cell Text: For plain text within cells
    - Calibri 10.5 font, No Spacing

    Cell Text Center: For plain text within cells that is center
    - Based on Cell Text
    - Alignment centered

    Cell Header: For header cells in the table
    - Based on Cell Text
    - Text is bolded

    Cell Header Right/Center: Both are for header cells with different alignments
    - Based on Header Text
    - Alignment Center or Right
    :return: Updated document
    """

    # TODO: Refactor this part into enum? separate module?
    styles = doc.styles

    style = styles.add_style('Cell Text', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['No Spacing']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.next_paragraph_style = styles['Normal']

    style = styles.add_style('Cell Text Center', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Cell Text']
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    style.next_paragraph_style = styles['Normal']

    style = styles.add_style('Cell Header', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Cell Text']
    style.font.bold = True

    style = styles.add_style('Cell Header Right', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Cell Header']
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    style = styles.add_style('Cell Header Center', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Cell Header']
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    style = styles.add_style('Header Footer Custom', WD_STYLE_TYPE.PARAGRAPH)
    doc_tab_stops = style.paragraph_format.tab_stops
    doc_tab_stops.add_tab_stop(Inches(tab_stops[0]), alignment=WD_TAB_ALIGNMENT.CENTER)
    doc_tab_stops.add_tab_stop(Inches(tab_stops[1]), alignment=WD_TAB_ALIGNMENT.RIGHT)
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    style = styles.add_style('Cell Heading', WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles['Heading 1']
    style.font.size = Pt(10.5)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.add_outline_level(1)


def __document_setup(template_name: str, template_desc: str, margins: tuple):
    """
    Sets up the initial format of the overall document
    :param template_name: Name of template document is based on
    :param template_desc: Description of template
    :param margins: margins for the document
    """
    # Margins
    doc.sections[0].left_margin = Inches(margins[0])
    doc.sections[0].right_margin = Inches(margins[1])
    doc.sections[0].top_margin = Inches(margins[2])
    doc.sections[0].bottom_margin = Inches(margins[3])

    # Headers
    doc.sections[0].header_distance = Inches(0.25)
    header_p = doc.sections[0].header.paragraphs[0]
    header_p.text = f"Excel Calculations Verification Tests\t{template_desc}\t{template_name}"
    header_p.style = doc.styles['Header Footer Custom']

    # Footers
    doc.sections[0].footer_distance = Inches(0.25)
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_r = footer_p.add_run('\t\tPage ')
    footer_r.add_field(__PAGE)
    footer_r.add_text(' of ')
    footer_r.add_field(__NUMPAGES)
    footer_p.style = doc.styles['Header Footer Custom']




