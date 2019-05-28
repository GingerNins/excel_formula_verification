import unittest
import re
from verification_document import TestDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt

DIRECTORY = r"C:/Users/erins/OneDrive - University of North Carolina at Chapel Hill/Programming Projects/" \
            r"Excel Formula Verification/"
TEMPLATE_DESC = r'Testing Document'
TEMPLATE_NAME = r'testing_document.xlsx'


class TestTestDocument(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.doc = TestDocument(DIRECTORY, TEMPLATE_NAME, TEMPLATE_DESC)

    def test_filename(self):
        filename_pattern_str = rf'{DIRECTORY}[0-9]{{8}}_ELS_{TEMPLATE_DESC}_[0-9]{{6}}\.docx'
        filename_pattern_re = re.compile(filename_pattern_str)
        self.assertTrue(filename_pattern_re.match(self.doc.filename))

        doc = self.doc._doc

    def test_styles(self):
        doc = self.doc._doc
        styles = doc.styles
        paragraph_styles = [s.name for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]

        self.assertTrue('Cell Text' in paragraph_styles)
        self.assertTrue('Cell Text Center' in paragraph_styles)
        self.assertTrue('Cell Header' in paragraph_styles)
        self.assertTrue('Cell Header Right' in paragraph_styles)
        self.assertTrue('Cell Header Center' in paragraph_styles)
        self.assertTrue('Header Footer Custom' in paragraph_styles)

        style = styles['Cell Text']
        self.assertEqual(styles['No Spacing'], style.base_style)
        self.assertEqual(styles['Normal'], style.next_paragraph_style)
        p = doc.add_paragraph('Cell Text Style Testing', style='Cell Text')
        self.assertEqual('Calibri', p.style.font.name)
        self.assertEqual(Pt(10.5), p.style.font.size)

        # Fonts and next paragraph styles are inherited from the base style
        # Normal <-- Cell Text <-- Cell Text Center, Cell Header <-- Cell Header Center/Right
        style = styles['Cell Text Center']
        self.assertEqual(styles['Cell Text'], style.base_style)
        self.assertEqual(styles['Normal'], style.next_paragraph_style)
        self.assertEqual(WD_PARAGRAPH_ALIGNMENT.CENTER, style.paragraph_format.alignment)

        style = styles['Cell Header']
        self.assertEqual(styles['Cell Text'], style.base_style)
        self.assertTrue(style.font.bold)

        style = styles['Cell Header Right']
        self.assertEqual(styles['Cell Header'], style.base_style)
        self.assertEqual(WD_PARAGRAPH_ALIGNMENT.RIGHT, style.paragraph_format.alignment)

        style = styles['Cell Header Center']
        self.assertEqual(style.base_style, styles['Cell Header'])
        self.assertEqual(style.paragraph_format.alignment, WD_PARAGRAPH_ALIGNMENT.CENTER)

        style = styles['Header Footer Custom']
        self.assertEqual(style.base_style, None)
        self.assertEqual('Calibri', style.font.name)
        self.assertEqual(Pt(10.5), style.font.size)
        tab_stops = style.paragraph_format.tab_stops
        self.assertEqual(2, len(tab_stops))
        self.assertEqual(Inches(4), tab_stops[0].position)
        self.assertEqual(Inches(7.5), tab_stops[1].position)
        self.assertEqual(WD_TAB_ALIGNMENT.CENTER, tab_stops[0].alignment)
        self.assertEqual(WD_TAB_ALIGNMENT.RIGHT, tab_stops[1].alignment)

    def test_setup(self):
        """
        Also incorporates the _add_field method
        """
        doc = self.doc._doc
        footer_string = f'\t\tPage Right-click to update field. of Right-click to update field.'
        self.assertEqual(footer_string, doc.sections[0].footer.paragraphs[0].text)
        header_string = f"Excel Calculations Verification Tests\t{TEMPLATE_DESC}\t{TEMPLATE_NAME}"
        self.assertEqual(header_string, doc.sections[0].header.paragraphs[0].text)

        self.assertEqual(Inches(0.5), doc.sections[0].left_margin)
        self.assertEqual(Inches(0.5), doc.sections[0].right_margin)
        self.assertEqual(Inches(0.5), doc.sections[0].top_margin)
        self.assertEqual(Inches(0.5), doc.sections[0].bottom_margin)

    def test_add_table(self):
        doc = self.doc._doc
        doc.add_table(self, 'f')


if __name__ == '__main__':
    unittest.main()
